from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "brochure_cliente"
OUT_DIR.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT_DIR / "Brochure_Profe360_Profesores.docx"
PDF_PATH = OUT_DIR / "Brochure_Profe360_Profesores.pdf"
LOGO_PATH = ROOT / "frontend" / "public" / "logo.png"
WEB_URL = "https://profe360cr.com/"
WA_URL = "https://api.whatsapp.com/send?phone=50686435071"
WA_LABEL = "WhatsApp: +506 8643-5071"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, name="Arial", size=11, bold=False, color="000000", italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    paragraph.alignment = align


def add_hyperlink(paragraph, url: str, text: str, color: str = "0563C1", bold: bool = False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    r_pr.append(color_el)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    if bold:
        bold_el = OxmlElement("w:b")
        r_pr.append(bold_el)

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(fonts)

    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, before=0, after=4, line=1.08)
    r = p.add_run(text)
    set_font(r, size=10.5, color="1F2937")


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.header_distance = Inches(0.3)
section.footer_distance = Inches(0.3)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)

header_table = doc.add_table(rows=1, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
header_table.autofit = False
header_table.columns[0].width = Inches(1.8)
header_table.columns[1].width = Inches(5.3)
left, right = header_table.rows[0].cells
left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
set_cell_margins(left, 70, 40, 70, 40)
set_cell_margins(right, 70, 120, 70, 120)

if LOGO_PATH.exists():
    p_logo = left.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_logo = p_logo.add_run()
    run_logo.add_picture(str(LOGO_PATH), width=Inches(1.12))

set_cell_shading(right, "0F2E4D")
p_contact = right.paragraphs[0]
style_paragraph(p_contact, before=0, after=0, line=1.05, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_hyperlink(p_contact, WEB_URL, WEB_URL, color="FFFFFF", bold=True)
p_contact.add_run("\n")
add_hyperlink(p_contact, WA_URL, WA_LABEL, color="DDEBFF")

p = doc.add_paragraph()
style_paragraph(p, before=10, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("PROFE360")
set_font(r, size=23, bold=True, color="0F2E4D")

p = doc.add_paragraph()
style_paragraph(p, before=0, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("Tecnologia inteligente para una docencia mas agil, organizada y efectiva")
set_font(r, size=13, color="1F4D78")

p = doc.add_paragraph()
style_paragraph(p, before=0, after=10, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("Menos tiempo en tareas administrativas, mas tiempo para ensenar.")
set_font(r, size=11.5, bold=True, color="B7791F")

p = doc.add_paragraph()
style_paragraph(p, before=0, after=8, line=1.1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
r = p.add_run(
    "Profe360 es una plataforma creada para transformar la experiencia del profesor, "
    "integrando en un solo sistema los procesos academicos, administrativos, de seguimiento "
    "y comunicacion que forman parte de su labor diaria."
)
set_font(r, size=10.5, color="1F2937")

p = doc.add_paragraph()
style_paragraph(p, before=0, after=8, line=1.1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
r = p.add_run(
    "Su valor esta en simplificar el trabajo docente, reducir la carga operativa y brindar "
    "herramientas que permiten actuar con mayor rapidez, mas control y mejor capacidad de respuesta."
)
set_font(r, size=10.5, color="1F2937")

p = doc.add_paragraph()
style_paragraph(p, before=4, after=6)
r = p.add_run("Beneficios para el profesor")
set_font(r, size=12.5, bold=True, color="0F2E4D")

for bullet in [
    "Centraliza asistencia, calificaciones, observaciones, reportes y seguimiento diario en un solo lugar.",
    "Reduce el tiempo invertido en tareas manuales y repetitivas.",
    "Facilita la comunicacion con encargados por medio de reportes y notificaciones por WhatsApp.",
    "Brinda mayor control, orden y respaldo en la gestion docente.",
    "Refuerza el Apoyo Educativo mediante el seguimiento de adecuaciones y una atencion mas oportuna a las necesidades del estudiante.",
    "Fortalece el seguimiento individual del estudiante y mejora la capacidad de respuesta del profesor.",
]:
    add_bullet(doc, bullet)

p = doc.add_paragraph()
style_paragraph(p, before=6, after=5)
r = p.add_run("Funciones destacadas")
set_font(r, size=12.5, bold=True, color="0F2E4D")

p = doc.add_paragraph()
style_paragraph(p, before=0, after=8, line=1.12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
r = p.add_run(
    "Asistencia y seguimiento diario, registro de notas, reportes en linea, boletas, certificaciones, "
    "Apoyo Educativo con seguimiento de adecuaciones, generacion de examenes con IA, tablas de especificaciones con IA, "
    "planeamientos con IA y Margarita, la asistente inteligente en tiempo real dentro de la plataforma."
)
set_font(r, size=10.5, color="1F2937")

callout = doc.add_table(rows=1, cols=1)
callout.alignment = WD_TABLE_ALIGNMENT.CENTER
callout.autofit = False
callout.columns[0].width = Inches(6.35)
cell = callout.rows[0].cells[0]
set_cell_shading(cell, "EAF3FF")
set_cell_margins(cell, 130, 180, 130, 180)
p = cell.paragraphs[0]
style_paragraph(p, before=0, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("Propuesta de valor")
set_font(r, size=12, bold=True, color="0F2E4D")
p = cell.add_paragraph()
style_paragraph(p, before=0, after=4, line=1.08, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run(
    "Profe360 permite que el profesor dedique menos esfuerzo a lo operativo y mas energia a lo pedagogico, "
    "fortaleciendo su gestion con organizacion, automatizacion, seguimiento e inteligencia aplicada a la educacion."
)
set_font(r, size=10.5, color="1F2937")
p = cell.add_paragraph()
style_paragraph(p, before=2, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
r = p.add_run("Con Profe360, el profesor gana tiempo, control y respaldo.")
set_font(r, size=11.5, bold=True, color="0F2E4D")

footer = section.footer
fp = footer.paragraphs[0]
style_paragraph(fp, before=0, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
add_hyperlink(fp, WEB_URL, "profe360cr.com", color="6B7280")
r = fp.add_run("  |  ")
set_font(r, size=9.5, color="6B7280")
add_hyperlink(fp, WA_URL, WA_LABEL, color="6B7280")

doc.save(DOCX_PATH)


styles = getSampleStyleSheet()
title_style = ParagraphStyle(
    "TitleCustom",
    parent=styles["Normal"],
    fontName="Helvetica-Bold",
    fontSize=22,
    leading=24,
    textColor=colors.HexColor("#0F2E4D"),
    alignment=1,
    spaceAfter=4,
)
subtitle_style = ParagraphStyle(
    "SubtitleCustom",
    parent=styles["Normal"],
    fontName="Helvetica",
    fontSize=12.5,
    leading=15,
    textColor=colors.HexColor("#1F4D78"),
    alignment=1,
    spaceAfter=4,
)
tagline_style = ParagraphStyle(
    "TaglineCustom",
    parent=styles["Normal"],
    fontName="Helvetica-Bold",
    fontSize=11.5,
    leading=14,
    textColor=colors.HexColor("#B7791F"),
    alignment=1,
    spaceAfter=10,
)
body_style = ParagraphStyle(
    "BodyCustom",
    parent=styles["Normal"],
    fontName="Helvetica",
    fontSize=10,
    leading=12.2,
    textColor=colors.HexColor("#1F2937"),
    alignment=4,
    spaceAfter=7,
)
section_style = ParagraphStyle(
    "SectionCustom",
    parent=styles["Normal"],
    fontName="Helvetica-Bold",
    fontSize=12,
    leading=14,
    textColor=colors.HexColor("#0F2E4D"),
    spaceAfter=4,
    spaceBefore=2,
)
bullet_style = ParagraphStyle(
    "BulletCustom",
    parent=styles["Normal"],
    fontName="Helvetica",
    fontSize=9.8,
    leading=11.6,
    leftIndent=14,
    firstLineIndent=-8,
    bulletIndent=0,
    textColor=colors.HexColor("#1F2937"),
    spaceAfter=3,
)
callout_title_style = ParagraphStyle(
    "CalloutTitle",
    parent=styles["Normal"],
    fontName="Helvetica-Bold",
    fontSize=11.5,
    leading=13,
    alignment=1,
    textColor=colors.HexColor("#0F2E4D"),
    spaceAfter=4,
)
callout_body_style = ParagraphStyle(
    "CalloutBody",
    parent=styles["Normal"],
    fontName="Helvetica",
    fontSize=9.9,
    leading=11.8,
    alignment=1,
    textColor=colors.HexColor("#1F2937"),
    spaceAfter=4,
)
callout_footer_style = ParagraphStyle(
    "CalloutFooter",
    parent=styles["Normal"],
    fontName="Helvetica-Bold",
    fontSize=10.6,
    leading=12.5,
    alignment=1,
    textColor=colors.HexColor("#0F2E4D"),
)
footer_style = ParagraphStyle(
    "FooterCustom",
    parent=styles["Normal"],
    fontName="Helvetica",
    fontSize=9.3,
    leading=11,
    alignment=1,
    textColor=colors.HexColor("#6B7280"),
)
contact_style = ParagraphStyle(
    "ContactStyle",
    parent=styles["Normal"],
    fontName="Helvetica",
    fontSize=10,
    leading=12,
    alignment=2,
    textColor=colors.white,
)

story = []

logo = Image(str(LOGO_PATH), width=0.98 * inch, height=0.98 * inch)
contact_box = Table(
    [[Paragraph(
        f"<link href='{WEB_URL}' color='white'><b>{WEB_URL}</b></link><br/>"
        f"<link href='{WA_URL}' color='#DDEBFF'>{WA_LABEL}</link>",
        contact_style,
    )]],
    colWidths=[4.0 * inch],
)
contact_box.setStyle(TableStyle([
    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#0F2E4D")),
    ("LEFTPADDING", (0, 0), (-1, -1), 12),
    ("RIGHTPADDING", (0, 0), (-1, -1), 12),
    ("TOPPADDING", (0, 0), (-1, -1), 8),
    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
]))
header = Table([[logo, contact_box]], colWidths=[1.75 * inch, 4.45 * inch])
header.setStyle(TableStyle([
    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ("LEFTPADDING", (0, 0), (-1, -1), 0),
    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
    ("TOPPADDING", (0, 0), (-1, -1), 0),
    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
]))
story.append(header)
story.append(Spacer(1, 0.12 * inch))
story.append(Paragraph("PROFE360", title_style))
story.append(Paragraph("Tecnologia inteligente para una docencia mas agil, organizada y efectiva", subtitle_style))
story.append(Paragraph("Menos tiempo en tareas administrativas, mas tiempo para ensenar.", tagline_style))
story.append(Paragraph(
    "Profe360 es una plataforma creada para transformar la experiencia del profesor, integrando en un solo sistema los procesos academicos, administrativos, de seguimiento y comunicacion que forman parte de su labor diaria.",
    body_style,
))
story.append(Paragraph(
    "Su valor esta en simplificar el trabajo docente, reducir la carga operativa y brindar herramientas que permiten actuar con mayor rapidez, mas control y mejor capacidad de respuesta.",
    body_style,
))
story.append(Paragraph("Beneficios para el profesor", section_style))
for item in [
    "Centraliza asistencia, calificaciones, observaciones, reportes y seguimiento diario en un solo lugar.",
    "Reduce el tiempo invertido en tareas manuales y repetitivas.",
    "Facilita la comunicacion con encargados por medio de reportes y notificaciones por WhatsApp.",
    "Brinda mayor control, orden y respaldo en la gestion docente.",
    "Refuerza el Apoyo Educativo mediante el seguimiento de adecuaciones y una atencion mas oportuna a las necesidades del estudiante.",
    "Fortalece el seguimiento individual del estudiante y mejora la capacidad de respuesta del profesor.",
]:
    story.append(Paragraph(item, bullet_style, bulletText="•"))
story.append(Paragraph("Funciones destacadas", section_style))
story.append(Paragraph(
    "Asistencia y seguimiento diario, registro de notas, reportes en linea, boletas, certificaciones, Apoyo Educativo con seguimiento de adecuaciones, generacion de examenes con IA, tablas de especificaciones con IA, planeamientos con IA y Margarita, la asistente inteligente en tiempo real dentro de la plataforma.",
    body_style,
))

callout = Table(
    [
        [Paragraph("Propuesta de valor", callout_title_style)],
        [Paragraph(
            "Profe360 permite que el profesor dedique menos esfuerzo a lo operativo y mas energia a lo pedagogico, fortaleciendo su gestion con organizacion, automatizacion, seguimiento e inteligencia aplicada a la educacion.",
            callout_body_style,
        )],
        [Paragraph("Con Profe360, el profesor gana tiempo, control y respaldo.", callout_footer_style)],
    ],
    colWidths=[6.2 * inch],
)
callout.setStyle(TableStyle([
    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EAF3FF")),
    ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#C9DCF5")),
    ("TOPPADDING", (0, 0), (-1, -1), 8),
    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ("LEFTPADDING", (0, 0), (-1, -1), 14),
    ("RIGHTPADDING", (0, 0), (-1, -1), 14),
]))
story.append(callout)
story.append(Spacer(1, 0.12 * inch))
story.append(Paragraph(
    f"<link href='{WEB_URL}' color='#6B7280'>profe360cr.com</link>  |  "
    f"<link href='{WA_URL}' color='#6B7280'>{WA_LABEL}</link>",
    footer_style,
))

pdf = SimpleDocTemplate(
    str(PDF_PATH),
    pagesize=letter,
    leftMargin=0.7 * inch,
    rightMargin=0.7 * inch,
    topMargin=0.65 * inch,
    bottomMargin=0.55 * inch,
)
pdf.build(story)

print(DOCX_PATH)
print(PDF_PATH)
